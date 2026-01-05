import re, os
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.shared import Cm
import streamlit as st

# 页面配置
st.set_page_config(
    page_title="Word一键排版工具",
    page_icon="📝",
    layout="wide"
)

# ========== 主页面：简洁的文件处理界面 ==========
st.title("📝 Word一键排版工具")
st.markdown("---")

# 简介
st.markdown("""
**一键智能排版，无需复杂设置！**

**功能特点：**
- 🎯 **智能识别**：自动识别文档标题层级和结构
- 🔢 **自动编号**：智能添加多级标题序号（中文数字方案）
- 🎨 **专业格式**：应用预设的专业排版格式
- ⚡ **批量处理**：支持多个文件同时处理
- 📥 **即传即用**：上传后立即处理，无需额外设置

**预设格式方案：**
- 📌 **标题格式**：1-3级标题自动编号，黑体/宋体字体
- 📝 **正文格式**：宋体/Times New Roman，10.5pt，首行缩进
- 📊 **表格格式**：统一字体，自动调整宽度
""")

# 文件上传区域
st.markdown("### 📤 上传Word文档")
uploaded_files = st.file_uploader(
    "选择Word文档 (.docx) - 支持多选",
    type=["docx"],
    accept_multiple_files=True,
    help="支持批量上传多个文档",
    label_visibility="collapsed"
)

# ========== 预设格式参数 ==========
# 这些是预设的格式规则，用户无需设置
PRESET_STYLES = {
    # 正文样式
    "body": {
        "cz_font_name": "宋体",
        "font_name": "Times New Roman",
        "font_size": 10.5,
        "space_before": 6.0,
        "space_after": 6.0,
        "line_spacing": 1.0,
        "first_line_indent": 0.75,  # cm
    },
    
    # 表格样式
    "table": {
        "cz_font_name": "宋体",
        "font_name": "Times New Roman",
        "font_size": 10.5,
        "space_before": 4.0,
        "space_after": 4.0,
        "line_spacing": 1.0,
        "width": 6.0,  # 英寸
    },
    
    # 标题样式 (1-3级)
    1: {
        'cz_font_name': '黑体',
        'font_name': 'Arial',
        'font_size': 14,
        'bold': False,
        'space_before': 12,
        'space_after': 12,
        'line_spacing': 1.5,
        'first_line_indent': 0,
    },
    2: {
        'cz_font_name': '黑体',
        'font_name': 'Arial',
        'font_size': 12,
        'bold': False,
        'space_before': 12,
        'space_after': 12,
        'line_spacing': 1.5,
        'first_line_indent': 0.75,
    },
    3: {
        'cz_font_name': '宋体',
        'font_name': 'Times New Roman',
        'font_size': 10.5,
        'bold': True,
        'space_before': 8,
        'space_after': 8,
        'line_spacing': 1.0,
        'first_line_indent': 1.5,
    },
}

# 使用中文数字编号方案
NUMBERING_SCHEME = "方案一：中文数字"

# 显示已上传文件
if uploaded_files:
    st.success(f"✅ 已选择 {len(uploaded_files)} 个文档")
    
    # 文件列表
    with st.expander("📋 文件列表", expanded=True):
        for i, file in enumerate(uploaded_files, 1):
            col1, col2, col3 = st.columns([6, 2, 2])
            with col1:
                st.write(f"**{file.name}**")
            with col2:
                st.write(f"`{file.size / 1024:.1f} KB`")
            with col3:
                st.write("📄")
    
    # 处理按钮
    st.markdown("---")
    
    # 处理按钮
    if st.button("🚀 一键智能排版", type="primary", use_container_width=True):
        # 创建进度条
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # 处理结果区域
        results_container = st.container()
        
        # 处理每个文件
        with results_container:
            for idx, uploaded_file in enumerate(uploaded_files):
                # 更新进度
                progress = (idx + 1) / len(uploaded_files)
                progress_bar.progress(progress)
                status_text.text(f"正在处理: **{uploaded_file.name}** ({idx + 1}/{len(uploaded_files)})")
                
                try:
                    # 处理文档
                    processed_buffer = process_single_document(
                        uploaded_file.read()
                    )
                    
                    # 显示处理结果
                    col_result1, col_result2 = st.columns([8, 2])
                    with col_result1:
                        st.write(f"✅ **{uploaded_file.name}** - 排版完成")
                    with col_result2:
                        st.download_button(
                            label="📥 下载文件",
                            data=processed_buffer,
                            file_name=f"排版_{uploaded_file.name}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_{idx}",
                            use_container_width=True
                        )
                    
                except Exception as e:
                    st.error(f"❌ 处理 {uploaded_file.name} 时出错: `{str(e)}`")
            
            # 完成提示
            progress_bar.empty()
            status_text.success("✅ 所有文档处理完成！")
            st.balloons()
            st.info("💡 所有文档已应用专业排版格式，标题已自动编号！")

else:
    st.info("📤 请上传需要排版的Word文档")

# ========== 工具函数定义 ==========
KNOWN_STYLES = {
    "Normal",
    "List Paragraph",
    "Heading 1", "Heading 2", "Heading 3", "Heading 4",
    "Heading 5", "Heading 6", "Heading 7", "Heading 8", "Heading 9"
}

def get_outline_level_from_xml(p):
    """从段落的XML中提取大纲级别，并加1"""
    xml = p._p.xml
    m = re.search(r'<w:outlineLvl w:val="(\d)"/>', xml)
    level = int(m.group(1)) if m else None
    if level is not None:
        level += 1
    return level

def restructure_outline(doc):
    """重构文档大纲"""
    for p in doc.paragraphs:
        zero_indent(p)
        lvl = get_outline_level_from_xml(p)
        if lvl and p.style.name == "Normal":
            heading_style = f"Heading {lvl}"
            if heading_style in doc.styles:
                p.style = doc.styles[heading_style]
    
    # 降级空标题
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and not p.text.strip():
            p.style = doc.styles["Normal"]

def zero_indent(p):
    """清除段落缩进"""
    pf = p.paragraph_format
    pf.left_indent = Cm(0)
    pf.first_line_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.tab_stops.clear_all()
    if p.text:
        p.text = p.text.lstrip()

def kill_all_numbering(doc):
    """清除所有编号"""
    for st_name in ['List Paragraph', 'Heading 1', 'Heading 2', 'Heading 3',
                    'Heading 4', 'Heading 5', 'Heading 6', 'Heading 7',
                    'Heading 8', 'Heading 9']:
        try:
            style = doc.styles[st_name]
        except KeyError:
            continue
        style_el = style._element
        for num_id in style_el.xpath('.//w:numId'):
            num_id.getparent().remove(num_id)

def set_font(run, cz_font_name, font_name):
    """设置字体"""
    rPr = run.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), cz_font_name)
    rFonts.set(qn('w:ascii'), font_name)

def num_to_cn(num):
    """数字转中文大写数字"""
    cn_nums = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    
    if num <= 10:
        return cn_nums[num]
    elif num < 20:
        return "十" + (cn_nums[num - 10] if num != 10 else "")
    elif num < 100:
        tens = num // 10
        ones = num % 10
        if ones == 0:
            return cn_nums[tens] + "十"
        else:
            return cn_nums[tens] + "十" + cn_nums[ones]
    else:
        return str(num)

def add_heading_numbers_custom(doc):
    """添加自定义标题序号（使用预设的中文数字方案）"""
    number_pattern = re.compile(
        r'^\s*'
        r'[（(]?'
        r'[\d一二三四五六七八九十零①②③④⑤⑥⑦⑧⑨⑩]{1,4}'
        r'[\.、）)\s]'
        r'(?:[\d一二三四五六七八九十零①②③④⑤⑥⑦⑧⑨⑩]{1,4}'
        r'[\.、）)\s]'
        r')*',
        re.UNICODE
    )
    
    heading_numbers = [0] * 9
    
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            if paragraph.text == "Ellipsis" or not paragraph.text.strip():
                continue
            
            # 清除原有编号
            paragraph.text = number_pattern.sub('', paragraph.text).strip()
            level = int(paragraph.style.name.split(' ')[1]) - 1
            
            # 更新序号
            heading_numbers[level] += 1
            for i in range(level + 1, len(heading_numbers)):
                heading_numbers[i] = 0
            
            # 添加序号（只处理1-3级标题）
            if heading_numbers[level] > 0 and level < 3:
                if level == 0:
                    # 一级标题：一、
                    number_str = num_to_cn(heading_numbers[0]) + "、"
                elif level == 1:
                    # 二级标题：（一）
                    number_str = "（" + num_to_cn(heading_numbers[1]) + "）"
                elif level == 2:
                    # 三级标题：1.
                    number_str = str(heading_numbers[2]) + "."
                else:
                    # 4级及以上标题：数字序号
                    number_str = str(heading_numbers[level]) + "."
                
                paragraph.text = number_str + paragraph.text

def process_single_document(file_bytes):
    """处理单个文档"""
    doc = Document(BytesIO(file_bytes))
    
    # 重构大纲
    restructure_outline(doc)
    
    # 清除编号
    kill_all_numbering(doc)
    
    # 添加标题序号
    add_heading_numbers_custom(doc)
    
    # 应用预设格式
    skipped = set()
    
    for p in doc.paragraphs:
        style_name = p.style.name
        
        if p.text == "Ellipsis" or not p.text.strip():
            continue
        
        if style_name not in KNOWN_STYLES:
            skipped.add(style_name)
            continue
        
        if style_name.startswith("Heading"):
            level = int(style_name.split(' ')[1])
            if level in PRESET_STYLES:
                rule = PRESET_STYLES[level]
                p.style.paragraph_format.space_before = Pt(rule['space_before'])
                p.style.paragraph_format.space_after = Pt(rule['space_after'])
                p.style.paragraph_format.line_spacing = rule['line_spacing']
                p.style.paragraph_format.first_line_indent = Cm(rule['first_line_indent'])
                for run in p.runs:
                    set_font(run, rule['cz_font_name'], rule['font_name'])
                    run.font.size = Pt(rule['font_size'])
                    run.font.bold = rule['bold']
        else:
            # 正文格式
            body_rule = PRESET_STYLES["body"]
            p.paragraph_format.space_before = Pt(body_rule['space_before'])
            p.paragraph_format.space_after = Pt(body_rule['space_after'])
            p.paragraph_format.line_spacing = body_rule['line_spacing']
            p.paragraph_format.first_line_indent = Cm(body_rule['first_line_indent'])
            for run in p.runs:
                set_font(run, body_rule['cz_font_name'], body_rule['font_name'])
                run.font.size = Pt(body_rule['font_size'])
    
    # 表格格式
    for tbl in doc.tables:
        tbl.width = Inches(PRESET_STYLES["table"]["width"])
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.style.name != "Normal":
                        skipped.add(f"表格内：{p.style.name}")
                        continue
                    for run in p.runs:
                        set_font(run, PRESET_STYLES["table"]["cz_font_name"], 
                                PRESET_STYLES["table"]["font_name"])
                        run.font.size = Pt(PRESET_STYLES["table"]["font_size"])
                    p.paragraph_format.space_before = Pt(PRESET_STYLES["table"]["space_before"])
                    p.paragraph_format.space_after = Pt(PRESET_STYLES["table"]["space_after"])
                    p.paragraph_format.line_spacing = PRESET_STYLES["table"]["line_spacing"]
    
    # 保存到buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 页脚
st.markdown("---")
st.caption("© 2024 Word一键排版工具 | 专业排版 • 简单易用")
