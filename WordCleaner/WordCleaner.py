import re
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.shared import Inches


# 标题样式
style_rules = {
    1: {'style_name': 'Heading 1', 'font_name': 'Arial','cz_font_name': '楷体', 'font_size': 10, 'bold': True, 'space_before': 12, 'space_after': 12, 'line_spacing': 1.5, 'first_line_indent': 18},
    2: {'style_name': 'Heading 2', 'font_name': 'Arial','cz_font_name': '宋体', 'font_size': 14, 'bold': True, 'space_before': 10, 'space_after': 10, 'line_spacing': 1.5, 'first_line_indent': 18},
    3: {'style_name': 'Heading 3', 'font_name': 'Arial','cz_font_name': '宋体','font_size': 12, 'bold': False, 'space_before': 8, 'space_after': 8, 'line_spacing': 1.5, 'first_line_indent': 0},
    4: {'style_name': 'Heading 4', 'font_name': 'Arial','cz_font_name': '宋体', 'font_size': 11, 'bold': False, 'space_before': 6, 'space_after': 6, 'line_spacing': 1.5, 'first_line_indent': 0},
    5: {'style_name': 'Heading 5', 'font_name': 'Arial','cz_font_name': '宋体', 'font_size': 10, 'bold': False, 'space_before': 4, 'space_after': 4, 'line_spacing': 1.5, 'first_line_indent': 0},
    6: {'style_name': 'Heading 6', 'font_name': 'Arial','cz_font_name': '宋体', 'font_size': 9, 'bold': False, 'space_before': 2, 'space_after': 2, 'line_spacing': 1.5, 'first_line_indent': 0},
    7: {'style_name': 'Heading 7', 'font_name': 'Arial','cz_font_name': '宋体', 'font_size': 8, 'bold': False, 'space_before': 0, 'space_after': 0, 'line_spacing': 1.0, 'first_line_indent': 18},
    8: {'style_name': 'Heading 8', 'font_name': 'Arial','cz_font_name': '宋体', 'font_size': 7, 'bold': False, 'space_before': 0, 'space_after': 0, 'line_spacing': 1.0, 'first_line_indent': 18},
    9: {'style_name': 'Heading 9', 'font_name': 'Arial','cz_font_name': '宋体', 'font_size': 6, 'bold': False, 'space_before': 0, 'space_after': 0, 'line_spacing': 1.0, 'first_line_indent': 18},

}

# 正文格式
bdy_cz_font_name = "宋体"  # 字体
bdy_font_name = "Times New Roman"
bdy_font_size = Pt(12)  # 字号
bdy_space_before = Pt(12)  # 段前行距
bdy_space_after = Pt(12)  # 段后行距
bdy_line_spacing = 1.0  #行距
bdy_first_line_indent = Inches(0.5)  # 首行缩进

# 表格格式
tbl_cz_font_name = "宋体"  # 中文字体
tbl_font_name = "Times New Roman"  # 英文字体
tbl_font_size = Pt(10)  # 表格字号
tbl_space_before = Pt(6)  # 表格段前行距
tbl_space_after = Pt(6)  # 表格段后行距
tbl_width = Inches(6)

def get_outline_level_from_xml(p):
    """
    从段落的XML中提取大纲级别，并加1
    """
    xml = p._p.xml
    m = re.search(r'<w:outlineLvl w:val="(\d)"/>', xml)
    level = int(m.group(1)) if m else None
    if level is not None:
        level += 1  # 加1
    return level
               
def set_font(run, cz_font_name, font_name):
    """
    设置字体。

    :param run: 文本运行对象
    :param chinese_font_name: 中文字体名称
    :param english_font_name: 英文字体名称
    """
    # 获取或创建字体属性
    rPr = run.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    # 设置中文字体和英文字体
    rFonts.set(qn('w:eastAsia'), cz_font_name)
    rFonts.set(qn('w:ascii'), font_name)
    
# 手动实现数字到中文大写数字的转换
def number_to_chinese(number):
    if number < 0 or number > 100:
        raise ValueError("数字必须在0到100之间")
    
    chinese_numbers = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九"]
    chinese_units = ["", "十", "百"]
    
    if number < 10:
        return chinese_numbers[number]
    elif number < 20:
        return "十" + (chinese_numbers[number - 10] if number != 10 else "")
    elif number < 100:
        tens = number // 10
        ones = number % 10
        return chinese_numbers[tens] + "十" + (chinese_numbers[ones] if ones != 0 else "")
    else:
        return "一百"
   
# 添加标题序号并清洗原有序号
def add_heading_numbers(doc):
    # 初始化标题序号
    heading_numbers = [0, 0, 0, 0, 0, 0, 0, 0, 0]  # 假设最多有九级标题

    # 定义不同层级的序号格式
    def format_number(level, number):
        if level == 0:
            return f"{number_to_chinese(number)}、"  # 第一层级：一、二、三、
        elif level == 1:
            return f"（{number_to_chinese(number)}）"  # 第二层级：（一）（二）（三）
        elif level == 2:
            return f"{number}."  # 第三层级：1.2.3.
        elif level == 3:
            return f"（{number}）"  # 第四层级：（1）（2）（3）
        elif level == 4:
            return f"{number}."  # 第五层级：1.2.3.
        elif level == 5:
            return f"（{number}）"  # 第六层级：（1）（2）（3）
        elif level == 6:
            return f"{number}."  # 第七层级：1.2.3.
        elif level == 7:
            return f"（{number}）"  # 第八层级：（1）（2）（3）
        elif level == 8:
            return f"{number}."  # 第九层级：1.2.3.
        else:
            return f"{number}."  # 默认格式

    # 定义正则表达式，匹配常见的序号格式
    number_pattern = re.compile(r'^[\d一二三四五六七八九十（）\.、\s]+')

    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        # 检查段落是否是标题
        if paragraph.style.name.startswith('Heading'):
            # 获取标题级别
            level = int(paragraph.style.name.split(' ')[1]) - 1

            # 清洗原文档中的序号
            paragraph.text = number_pattern.sub('', paragraph.text).strip()

            # 更新序号
            heading_numbers[level] += 1
            for i in range(level + 1, len(heading_numbers)):
                heading_numbers[i] = 0  # 重置下级标题序号

            # 构造序号字符串
            number_str = format_number(level, heading_numbers[level])

            # 添加序号到标题文本
            paragraph.text = number_str + paragraph.text

def modify_document_format(doc):
    """
    修改 Word 文档中正文和表格的格式。

    :param file_path: 输入的 Word 文档路径
    :param output_path: 输出的 Word 文档路径，默认为 "modified.docx"
    """    
    # 遍历文档中的每个段落
    for paragraph in doc.paragraphs:
        # 检查是否是标题（标题的 style 通常以 "Heading" 开头）
        if  paragraph.style.name.startswith("Heading"):
            style_name = paragraph.style.name
            # 查找匹配的样式规则
            for level, rule in style_rules.items():
                if rule['style_name'] == style_name:
                    # 修改段前段后行距和首行缩进
                    paragraph.style.paragraph_format.space_before = Pt(rule['space_before'])
                    paragraph.style.paragraph_format.space_after = Pt(rule['space_after'])
                    paragraph.style.paragraph_format.line_spacing = rule['line_spacing']
                    paragraph.style.paragraph_format.first_line_indent = Pt(rule['first_line_indent'])
                    # 修改字体字号和粗体
                    for run in paragraph.runs:
                        set_font(run, rule['cz_font_name'], rule['font_name'])
                        run.font.size = Pt(rule['font_size'])
                        run.font.bold = rule['bold']
        else:            
            # 修改段前段后行距和首行缩进
            paragraph.paragraph_format.space_before = bdy_space_before
            paragraph.paragraph_format.space_after = bdy_space_after
            paragraph.paragraph_format.line_spacing = bdy_line_spacing
            paragraph.paragraph_format.first_line_indent = bdy_first_line_indent
            # 修改字体字号
            for run in paragraph.runs:
                set_font(run, bdy_cz_font_name, bdy_font_name)
                run.font.size = bdy_font_size

                
    # 遍历文档中的每个表格
    for table in doc.tables:
        table.width = tbl_width 
        # 遍历表格中的每个单元格
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # 修改字体和字号
                    for run in paragraph.runs:
                        # 设置中文字体和英文字体
                        set_font(run, tbl_cz_font_name, tbl_font_name)
                        # 设置字号
                        run.font.size = tbl_font_size

                    # 修改段前段后行距
                    paragraph.paragraph_format.space_before = tbl_space_before
                    paragraph.paragraph_format.space_after = tbl_space_after
    
# 主程序
def main():
    # 获取 Python 所在文件夹路径
    current_folder = os.path.dirname(os.path.abspath(__file__))

    # 获取文件夹下所有 .docx 文件
    docx_files = [f for f in os.listdir(current_folder) if f.endswith('.docx')]

    if not docx_files:
        print(f"文件夹 {current_folder} 中没有找到任何 .docx 文件。")
        return

    for file_name in docx_files:
        file_path = os.path.join(current_folder, file_name)

        # 打开一个现有的 Word 文档
        doc = Document(file_path)
        
        for para in doc.paragraphs:
            outline_level = get_outline_level_from_xml(para)
            style_name = para.style.name

            # 如果获取到大纲级别且当前样式为正文，根据大纲级别设置对应的标题样式
            if outline_level is not None and style_name == 'Normal':
                # 根据大纲级别设置标题样式
                if outline_level == 1:
                    para.style = doc.styles['Heading 1']
                elif outline_level == 2:
                    para.style = doc.styles['Heading 2']
                elif outline_level == 3:
                    para.style = doc.styles['Heading 3']
                elif outline_level == 4:
                    para.style = doc.styles['Heading 4']
                elif outline_level == 5:
                    para.style = doc.styles['Heading 5']
                elif outline_level == 6:
                    para.style = doc.styles['Heading 6']
                elif outline_level == 7:
                    para.style = doc.styles['Heading 7']
                elif outline_level == 8:
                    para.style = doc.styles['Heading 8']
                elif outline_level == 9:
                    para.style = doc.styles['Heading 9']
            
        # 添加标题序号并清洗原有序号
        add_heading_numbers(doc)

        # 应用样式规则
        modify_document_format(doc)

        # 构造输出文件名
        output_file_name = f"{os.path.splitext(file_name)[0]}_已修改.docx"
        output_file_path = os.path.join(current_folder, output_file_name)

        # 保存修改后的文档
        doc.save(output_file_path)
        print(f"文件 {file_name} 已修改并保存为 {output_file_path}")

if __name__ == "__main__":
    main()
